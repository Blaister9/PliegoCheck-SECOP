"""Pruebas del runner y extractores deterministas."""

import hashlib
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

import fitz
from docx import Document as DocxDocument
from openpyxl import Workbook
from sqlalchemy import Select, func, select

from pliegocheck_api.config import get_settings
from pliegocheck_api.db import get_sessionmaker
from pliegocheck_api.models import (
    DocumentExtraction,
    DocumentProcessingJob,
    ExtractedSegment,
    Process,
    ProcessDocument,
)
from pliegocheck_api.storage import LocalDocumentStorage
from pliegocheck_schemas import (
    DocumentProcessingJobType,
    DocumentProcessingStatus,
    DocumentType,
    DocumentUploadStatus,
    ProcessSource,
    ProcessStatus,
)
from pliegocheck_worker.runner import claim_next_job, drain, run_once


def test_run_once_without_jobs_returns_idle() -> None:
    assert run_once("test-worker")["processed"] == 0


def test_worker_extracts_txt_with_line_ranges(tmp_path: Path) -> None:
    document_id = create_document(tmp_path, "notas.txt", b"Linea 1\nLinea 2\n")
    result = run_once("txt-worker")
    assert result["job_status"] == "COMPLETED"
    extraction, segments, document = read_result(document_id)
    assert extraction.status in {"COMPLETED", "COMPLETED_WITH_WARNINGS"}
    assert document.processing_status in {
        DocumentProcessingStatus.COMPLETED.value,
        DocumentProcessingStatus.COMPLETED_WITH_WARNINGS.value,
    }
    assert segments[0].segment_type == "TEXT_LINES"
    assert segments[0].line_start == 1
    assert "Linea 2" in segments[0].text


def test_worker_extracts_pdf_by_page(tmp_path: Path) -> None:
    path = tmp_path / "pliego.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Pagina uno trazable")
    page = pdf.new_page()
    page.insert_text((72, 72), "Pagina dos trazable")
    pdf.save(path)
    pdf.close()
    document_id = create_document(tmp_path, "pliego.pdf", path.read_bytes())
    result = run_once("pdf-worker")
    assert result["job_status"] == "COMPLETED"
    extraction, segments, _ = read_result(document_id)
    assert extraction.status in {"COMPLETED", "COMPLETED_WITH_WARNINGS"}
    assert extraction.page_count == 2
    assert [segment.page_number for segment in segments] == [1, 2]


def test_worker_marks_empty_pdf_as_needs_ocr(tmp_path: Path) -> None:
    path = tmp_path / "scanned.pdf"
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(path)
    pdf.close()
    document_id = create_document(tmp_path, "scanned.pdf", path.read_bytes())
    run_once("ocr-worker")
    extraction, _, document = read_result(document_id)
    assert extraction.status == "NEEDS_OCR"
    assert document.processing_status == DocumentProcessingStatus.NEEDS_OCR.value


def test_worker_marks_encrypted_pdf(tmp_path: Path) -> None:
    path = tmp_path / "encrypted.pdf"
    pdf = fitz.open()
    pdf.new_page().insert_text((72, 72), "Secreto")
    pdf.save(
        path,
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner",
        user_pw="user",
    )
    pdf.close()
    document_id = create_document(tmp_path, "encrypted.pdf", path.read_bytes())
    run_once("encrypted-worker")
    extraction, _, document = read_result(document_id)
    assert extraction.status == "ENCRYPTED"
    assert document.processing_status == DocumentProcessingStatus.ENCRYPTED.value


def test_worker_extracts_docx_paragraph_and_table(tmp_path: Path) -> None:
    path = tmp_path / "anexo.docx"
    doc = DocxDocument()
    doc.add_paragraph("Parrafo trazable")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    doc.save(str(path))
    document_id = create_document(tmp_path, "anexo.docx", path.read_bytes())
    run_once("docx-worker")
    extraction, segments, _ = read_result(document_id)
    assert extraction.status in {"COMPLETED", "COMPLETED_WITH_WARNINGS"}
    assert [segment.segment_type for segment in segments] == ["PARAGRAPH", "TABLE"]


def test_worker_extracts_xlsx_rows_without_evaluating_formulas(tmp_path: Path) -> None:
    path = tmp_path / "anexo.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "Hoja 1"
    sheet.append(["Concepto", "Valor"])
    sheet.append(["Formula", "=1+1"])
    workbook.save(path)
    document_id = create_document(tmp_path, "anexo.xlsx", path.read_bytes())
    run_once("xlsx-worker")
    extraction, segments, _ = read_result(document_id)
    assert extraction.status == "COMPLETED_WITH_WARNINGS"
    assert segments[0].segment_type == "SHEET_ROW"
    assert "=1+1" in segments[0].text
    assert extraction.sheet_count == 1


def test_worker_extracts_csv_rows(tmp_path: Path) -> None:
    document_id = create_document(tmp_path, "tabla.csv", b"nombre,valor\nuno,1\n")
    run_once("csv-worker")
    extraction, segments, _ = read_result(document_id)
    assert extraction.status in {"COMPLETED", "COMPLETED_WITH_WARNINGS"}
    assert segments[0].sheet_name == "csv"
    assert "nombre\tvalor" in segments[0].text


def test_worker_marks_images_and_legacy_formats(tmp_path: Path) -> None:
    image_id = create_document(tmp_path, "imagen.png", b"\x89PNG\r\n\x1a\npayload")
    legacy_id = create_document(tmp_path, "legacy.doc", b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
    summary = drain(max_jobs=10, worker_id="format-worker")
    assert summary["processed"] == 2
    image_extraction, _, image = read_result(image_id)
    legacy_extraction, _, legacy = read_result(legacy_id)
    assert image_extraction.status == "NEEDS_OCR"
    assert image.processing_status == DocumentProcessingStatus.NEEDS_OCR.value
    assert legacy_extraction.status == "UNSUPPORTED"
    assert legacy.processing_status == DocumentProcessingStatus.UNSUPPORTED.value


def test_worker_is_idempotent_for_same_hash_and_version(tmp_path: Path) -> None:
    document_id = create_document(tmp_path, "notas.txt", b"Texto idempotente\n")
    run_once("first-worker")
    with get_sessionmaker()() as session:
        document = session.get(ProcessDocument, document_id)
        assert document is not None
        session.add(
            DocumentProcessingJob(
                id=uuid4(),
                document_id=document_id,
                job_type=DocumentProcessingJobType.EXTRACT_DOCUMENT.value,
                max_attempts=3,
                available_at=datetime.now(UTC),
            )
        )
        document.processing_status = DocumentProcessingStatus.QUEUED.value
        session.commit()
    second = run_once("second-worker")
    assert second["idempotent"] is True
    with get_sessionmaker()() as session:
        assert session.scalar(select_count(DocumentExtraction)) == 1
        assert session.scalar(select_count(ExtractedSegment)) == 1


def test_two_claims_do_not_take_same_job(tmp_path: Path) -> None:
    create_document(tmp_path, "notas.txt", b"Trabajo unico\n")
    sessionmaker = get_sessionmaker()
    with sessionmaker() as first_session, sessionmaker() as second_session:
        first = claim_next_job(first_session, "worker-a")
        second = claim_next_job(second_session, "worker-b")
    assert first is not None
    assert second is None


def create_document(tmp_path: Path, filename: str, content: bytes) -> UUID:
    source = tmp_path / filename
    source.write_bytes(content)
    digest = hashlib.sha256(content).hexdigest()
    process_id = uuid4()
    document_id = uuid4()
    storage_key = f"{process_id}/{document_id}/{document_id.hex}{Path(filename).suffix}"
    storage = LocalDocumentStorage(get_settings().storage_path)
    storage.save(source, storage_key)
    with get_sessionmaker()() as session:
        process = Process(
            id=process_id,
            internal_reference=f"MAN-TEST-{process_id.hex[:8]}",
            title="Proceso worker",
            contracting_entity="Entidad",
            status=ProcessStatus.READY_FOR_INVENTORY.value,
            source=ProcessSource.MANUAL.value,
        )
        document = ProcessDocument(
            id=document_id,
            process_id=process_id,
            original_filename=filename,
            stored_filename=f"{document_id.hex}{Path(filename).suffix}",
            storage_key=storage_key,
            declared_content_type="application/octet-stream",
            detected_content_type="application/octet-stream",
            extension=Path(filename).suffix.lower(),
            size_bytes=len(content),
            sha256=digest,
            document_type=DocumentType.UNKNOWN.value,
            upload_status=DocumentUploadStatus.STORED.value,
            processing_status=DocumentProcessingStatus.QUEUED.value,
        )
        job = DocumentProcessingJob(
            id=uuid4(),
            document_id=document_id,
            job_type=DocumentProcessingJobType.EXTRACT_DOCUMENT.value,
            max_attempts=3,
            available_at=datetime.now(UTC),
        )
        session.add_all([process, document, job])
        session.commit()
    return document_id


def read_result(
    document_id: UUID,
) -> tuple[DocumentExtraction, list[ExtractedSegment], ProcessDocument]:
    with get_sessionmaker()() as session:
        extraction = session.scalars(
            select(DocumentExtraction).where(DocumentExtraction.document_id == document_id)
        ).one()
        segments = list(
            session.scalars(
                select(ExtractedSegment)
                .where(ExtractedSegment.extraction_id == extraction.id)
                .order_by(ExtractedSegment.sequence)
            ).all()
        )
        document = session.get(ProcessDocument, document_id)
        assert document is not None
        session.expunge(extraction)
        for segment in segments:
            session.expunge(segment)
        session.expunge(document)
        return extraction, segments, document


def select_count(model: type[Any]) -> Select[tuple[int]]:
    return select(func.count()).select_from(model)
