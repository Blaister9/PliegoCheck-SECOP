"""Extractor DOCX deterministico."""

from docx import Document

from pliegocheck_worker.extraction.limits import ExtractionLimits
from pliegocheck_worker.extraction.models import (
    ControlledExtractionError,
    ExtractionResultData,
    ExtractionWarningData,
    SegmentData,
)
from pliegocheck_worker.extraction.security import inspect_zip_container
from pliegocheck_worker.extraction.text import ensure_character_budget, normalize_text


def extract_docx(path: str, limits: ExtractionLimits) -> ExtractionResultData:
    inspect_zip_container(path, limits)
    try:
        document = Document(path)
    except Exception as exc:
        raise ControlledExtractionError(
            "EXTRACTION_FAILED", "No fue posible abrir el DOCX."
        ) from exc

    segments: list[SegmentData] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        text = normalize_text(paragraph.text)
        if text:
            segments.append(
                SegmentData(
                    sequence=len(segments) + 1,
                    segment_type="PARAGRAPH",
                    text=text,
                    paragraph_index=paragraph_index,
                    source_location={"paragraph_index": paragraph_index},
                )
            )

    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [normalize_text(cell.text).replace("\t", " ") for cell in row.cells]
            row_text = "\t".join(cells).strip()
            if row_text:
                rows.append(row_text)
        table_text = normalize_text("\n".join(rows))
        if table_text:
            segments.append(
                SegmentData(
                    sequence=len(segments) + 1,
                    segment_type="TABLE",
                    text=table_text,
                    table_index=table_index,
                    source_location={"table_index": table_index, "format": "tsv"},
                )
            )

    character_count = ensure_character_budget(
        (segment.text for segment in segments),
        limits.max_characters,
    )
    if not segments:
        return ExtractionResultData(
            status="COMPLETED_WITH_WARNINGS",
            detected_format="docx",
            warnings=[
                ExtractionWarningData(
                    code="DOCX_WITHOUT_TEXT",
                    message="El DOCX no contiene parrafos ni tablas con texto.",
                )
            ],
            character_count=0,
        )
    return ExtractionResultData(
        status="COMPLETED",
        detected_format="docx",
        segments=segments,
        character_count=character_count,
    )
